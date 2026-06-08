#!/usr/bin/env python3
"""Build Technical_Report.docx for the Formative 2 submission."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)


def H2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)


def P(text):
    doc.add_paragraph(text)


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Menlo"
    r.font.size = Pt(9)


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = str(val)


# ---------- Title ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Technical Report\nFormative 2: Blockchain Transaction Models and Mining Simulation")
r.bold = True
r.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("ALU Blockchain Attendance System")
sr.italic = True
sr.font.size = Pt(12)

doc.add_paragraph("")

# ---------- 1. Attendance to Transactions ----------
H1("1. Attendance to Transactions")
P(
    "When a student is marked, the system immediately constructs an "
    "attendance block holding the student's ID, name, course code (looked "
    "up from the registry), status, timestamp, the token reward for that "
    "status, and a tx_id equal to the SHA-256 hash of the reward "
    "transaction data."
)
table(
    ["Status", "Token reward", "Transaction generated"],
    [
        ["PRESENT", "10 coins", "yes"],
        ["LATE", "5 coins", "yes"],
        ["ABSENT", "0 coins", "no"],
    ],
)
P(
    "ABSENT records are still placed in the pending pool so the attendance "
    "is captured, but they carry token_reward = 0 and no balance update "
    "happens when they are confirmed."
)

# ---------- 2. Pending Pool and Mining ----------
H1("2. Pending Pool and Mining")
P(
    "Attendance records are NOT immediately permanent. They sit in a singly "
    "linked pending pool. Mining is the confirmation step that promotes a "
    "block from pending to confirmed:"
)
for line in [
    "1. The miner assigns the next chain index and copies the current chain tip into the block's previous_hash.",
    "2. The miner signs the block (HMAC-style SHA-256 over the signed fields plus a system secret).",
    "3. Proof of work increments a nonce and recomputes SHA-256(index | student_id | course | status | timestamp | reward | tx_id | previous_hash | signature | nonce) until the hash starts with the required number of leading 0 characters.",
    "4. The block's hash becomes the new chain tip; the block is appended to the confirmed list and the token transaction updates balances under the active model.",
]:
    doc.add_paragraph(line, style="List Number")
P(
    "Because the PoW input includes both previous_hash and signature, any "
    "change to a prior block invalidates that block's hash AND forces "
    "re-mining of every block that follows."
)

# ---------- 3. UTXO Implementation ----------
H1("3. UTXO Implementation")
P(
    "Each token credit (reward or transfer) appends a new node to a UTXO "
    "linked list. A UTXO holds an id (SHA-256 of owner:amount:counter), "
    "the owner, the amount, and a spent flag."
)
P(
    "Balance for a student is the sum of amount over all UTXOs owned by "
    "them where spent == 0."
)
P(
    "Reward credit: confirming an attendance block with reward R > 0 "
    "creates a UTXO of value R - 1 for the student (the fixed 1-coin fee "
    "is destroyed). The full UTXO set is printed after every confirmed "
    "block so the grader can watch the set grow."
)
P(
    "Spending: a transfer walks the sender's unspent UTXOs marking them "
    "spent until total >= amount + fee. If total < amount + fee, the "
    "transfer is rejected. If total > amount + fee, the excess is returned "
    "to the sender as a fresh change UTXO. The receiver gets a new UTXO "
    "worth exactly amount."
)
P(
    "Double-spend prevention is built in: once a UTXO is marked spent = 1, "
    "every subsequent transfer skips it."
)

# ---------- 4. Account model ----------
H1("4. Account-Based Implementation")
P(
    "Each student record holds a balance, a nonce (initially 0), and a "
    "linked-list history of every transaction touching that account."
)
P(
    "Reward credit: confirming a reward of R credits the student with "
    "R - 1 and prepends a history entry with sender SYSTEM."
)
P(
    "Manual transfer debits the sender by amount + fee, credits the "
    "receiver by amount, increments the sender's nonce, and prepends one "
    "entry to each side's history with sender, receiver, amount, fee, "
    "nonce."
)
P(
    "Validation: the transfer is rejected if (a) the sender's balance is "
    "less than amount + fee, or (b) the submitted nonce is not exactly "
    "current_nonce + 1. The nonce check prevents replayed transactions."
)
P("History per account is browsable from the menu (option 12) by student ID.")

# ---------- 5. Comparison ----------
H1("5. Model Comparison")
table(
    ["Property", "UTXO", "Account"],
    [
        ["Balance lookup", "O(n) scan of UTXO list", "O(1) field read"],
        ["Double-spend prevention", "Built into the spent flag on each UTXO", "Requires the nonce check"],
        ["State size", "Grows with credits ever issued", "Bounded by number of accounts"],
        ["Conceptual fit", "Cash-like; each coin is traceable", "Bank-account-like"],
        ["Implementation effort", "Higher (data structure + transactions)", "Lower (one field per student)"],
    ],
)

# ---------- 6. Mining Methodology ----------
H1("6. Mining Methodology")

H2("6.1 Solo Mining")
P(
    "One miner runs the proof-of-work loop on every pending block in order. "
    "The miner is paid the full BLOCK_REWARD = 25 per confirmed block on top "
    "of the per-student token rewards. The number of hash attempts before a "
    "valid hash was found is printed alongside each block."
)

H2("6.2 Pool Mining")
P(
    "Three simulated miners are each assigned a random hash rate. The pool "
    "mines each pending block. A pool_fee = BLOCK_REWARD * 2% = 0.50 is "
    "retained by the pool; the remaining 24.50 is distributed according to "
    "share_i = (attempts_i / sum_attempts) * 24.50."
)

H2("6.3 Cloud Mining")
P(
    "The user rents hash power for 1 to 5 rounds at a configurable rental "
    "fee per round. Each round credits CLOUD_REWARD = 18 and debits the "
    "rental fee. A warning is printed at any round where cumulative fees "
    "exceed cumulative rewards. After the rental period, the rented rig "
    "confirms the pending blocks via the same PoW loop."
)

H2("6.4 Difficulty")
P(
    "difficulty is the number of leading 0 characters required of the "
    "block's SHA-256 hex hash. It is configurable from 1 to 4 via the "
    "--difficulty CLI flag or menu option 2 (default 2). Higher difficulty "
    "multiplies expected attempts by roughly 16x per level."
)

# ---------- 7. Security ----------
H1("7. Security Mechanisms")
for line in [
    "SHA-256 everywhere: tx_id, signature, and hash are all SHA-256 digests of well-defined input strings.",
    "Chain linkage: every block carries previous_hash; option 13 walks the chain comparing each block's previous_hash against the preceding block's hash, and recomputes the SHA-256 commitment.",
    "Signature: each confirmed block is committed to a system secret via HMAC-style SHA-256 over the signed fields. verify_signature recomputes and compares.",
    "Proof of work: the PoW hash input includes previous_hash and signature, so tampering with any prior block invalidates its own hash AND forces re-mining of every subsequent block.",
    "Tamper demo: option 14 alters a confirmed block's status in memory without re-signing or re-mining; option 13 then prints TAMPERED messages and the change is reverted.",
]:
    doc.add_paragraph(line, style="List Bullet")

# ---------- 8. Screenshots ----------
H1("8. Screenshots (Terminal Captures)")

H2("8.1 Registry Load and Attendance Marking")
code(
    "Loaded 5 student(s) from students.txt\n"
    "Model=UTXO difficulty=2 genesis=0000000000000000...\n"
    "\n"
    "Registered students (5):\n"
    "  ALU001     John Doe                       BLK101\n"
    "  ALU002     Jane Smith                     BLK101\n"
    "  ALU003     Amara Diallo                   BLK101\n"
    "  ALU004     Kwame Mensah                   BLK101\n"
    "  ALU005     Fatima Bello                   BLK102\n"
    "\n"
    "> Mark attendance\n"
    "Student ID: ALU001\n"
    "Status (PRESENT/LATE/ABSENT): PRESENT\n"
    "Attendance placed in pending pool. tx=c8a4775e4d4a reward=10\n"
    "\n"
    "Student ID: ALU003\n"
    "Status (PRESENT/LATE/ABSENT): ABSENT\n"
    "Attendance placed in pending pool. tx=973fd57ad48e reward=0\n"
    "ABSENT recorded: no token transaction generated.\n"
    "\n"
    "Student ID: ALU999\n"
    "ERROR: Student ID not found"
)

H2("8.2 Pending Pool")
code(
    "Pending pool:\n"
    "  ALU001 BLK101 PRESENT reward=10 tx=c8a4775e4d4a\n"
    "  ALU002 BLK101 LATE    reward=5  tx=79d4eeb26811\n"
    "  ALU003 BLK101 ABSENT  reward=0  tx=973fd57ad48e"
)

H2("8.3 Solo Mining + Confirmation + UTXO After Each Block")
code(
    "Solo mining at difficulty 2 (hash must start with 2 zero(s))\n"
    "Solo mined block #0 tx=c8a4775e4d4a attempts=112 hash=00ec3526eb65...\n"
    "Solo mined block #1 tx=79d4eeb26811 attempts=91  hash=00f1c9b0dd9d...\n"
    "Solo mined block #2 tx=973fd57ad48e attempts=31  hash=00e37d0fe62b...\n"
    "Solo miner reward=25 total_attempts=234\n"
    "\n"
    "Confirmed by solo mining: block=0 John Doe BLK101 PRESENT reward=10 signature=VALID\n"
    "UTXO set:\n"
    "  930147f1651b owner=ALU001 amount=9 spent=no\n"
    "\n"
    "Confirmed by solo mining: block=1 Jane Smith BLK101 LATE reward=5 signature=VALID\n"
    "UTXO set:\n"
    "  1eada2f749d6 owner=ALU002 amount=4 spent=no\n"
    "  930147f1651b owner=ALU001 amount=9 spent=no\n"
    "\n"
    "Confirmed by solo mining: block=2 Amara Diallo BLK101 ABSENT reward=0 signature=VALID\n"
    "UTXO set:\n"
    "  1eada2f749d6 owner=ALU002 amount=4 spent=no\n"
    "  930147f1651b owner=ALU001 amount=9 spent=no"
)

H2("8.4 Pool Mining Reward Table")
code(
    "Pool mining (block reward=25, pool fee=2%=0.50, pool=24.50):\n"
    "Miner | Attempts | Share %  | Reward\n"
    "------+----------+----------+--------\n"
    "M1    | 240      |   19.98  |   4.90\n"
    "M2    | 61       |    5.08  |   1.24\n"
    "M3    | 900      |   74.94  |  18.36\n"
    "Total distributed=24.50  retained_by_pool=0.50"
)

H2("8.5 Cloud Mining (Unprofitable Edge Case)")
code(
    "Rental duration rounds (1-5): 3\n"
    "Rental fee per round (try 12 profitable, 22 unprofitable): 22\n"
    "Cloud reward per round=18, rental fee per round=22\n"
    "Round 1 gross=18 fees=22 net=-4\n"
    "Warning: cloud rental is unprofitable at round 1.\n"
    "Round 2 gross=36 fees=44 net=-8\n"
    "Warning: cloud rental is unprofitable at round 2.\n"
    "Round 3 gross=54 fees=66 net=-12\n"
    "Warning: cloud rental is unprofitable at round 3.\n"
    "Cloud mining summary gross=54 total_fees=66 net_profit=-12\n"
    "Cloud rig confirmed block #0 tx=f5f8fe45560a attempts=77\n"
    "Cloud miner total_attempts=77\n"
    "Confirmed by cloud mining: block=0 John Doe BLK101 PRESENT reward=10 signature=VALID"
)

H2("8.6 Balances and UTXO Set")
code(
    "Student balances using UTXO model:\n"
    "  ALU001 John Doe       balance=9 nonce=0\n"
    "  ALU002 Jane Smith     balance=4 nonce=0\n"
    "  ALU003 Amara Diallo   balance=0 nonce=0\n"
    "  ALU004 Kwame Mensah   balance=0 nonce=0\n"
    "  ALU005 Fatima Bello   balance=0 nonce=0\n"
    "\n"
    "UTXO set:\n"
    "  1eada2f749d6 owner=ALU002 amount=4 spent=no\n"
    "  930147f1651b owner=ALU001 amount=9 spent=no"
)

H2("8.7 Chain Validation and Tamper Detection")
code(
    "Validating confirmed chain...\n"
    "Chain is VALID. All hashes, signatures and links verified.\n"
    "\n"
    "> Tamper detection demo\n"
    "Block index to tamper with (0-2): 1\n"
    "Forged block 1 status: LATE -> ABSENT (hash/signature NOT updated)\n"
    "\n"
    "Validating confirmed chain...\n"
    "  [Block 1] TAMPERED: stored hash does not match contents\n"
    "  [Block 1] TAMPERED: signature does not verify\n"
    "Chain is INVALID: 2 problem(s) detected.\n"
    "In-memory change reverted."
)

H2("8.8 Account Model: Nonce Reject, Insufficient Balance, History")
code(
    "Choose model: 2\n"
    "Using account model.\n"
    "\n"
    "(mark ALU001 PRESENT twice, solo mine, balance=18 nonce=0)\n"
    "\n"
    "Manual transfer ALU001 -> ALU002 amount=5 nonce=5\n"
    "Rejected: incorrect or reused nonce.\n"
    "\n"
    "Manual transfer ALU001 -> ALU002 amount=100 nonce=1\n"
    "Rejected: insufficient account balance.\n"
    "\n"
    "Manual transfer ALU001 -> ALU002 amount=5 nonce=1\n"
    "Account transfer accepted. Sender nonce is now 1.\n"
    "\n"
    "Student balances using account model:\n"
    "  ALU001 John Doe   balance=12 nonce=1\n"
    "  ALU002 Jane Smith balance=5  nonce=0\n"
    "\n"
    "History for ALU001:\n"
    "  ALU001 -> ALU002 amount=5 fee=1 nonce=1\n"
    "  SYSTEM -> ALU001 amount=9 fee=1 nonce=0\n"
    "  SYSTEM -> ALU001 amount=9 fee=1 nonce=0"
)

# ---------- 9. Edge Case Test Results ----------
H1("9. Edge Case Test Results")
table(
    ["Case", "Trigger", "Observed output"],
    [
        ["Absent -> no transaction", "Mark ALU003 as ABSENT, mine, check balances", "Balance for ALU003 stays 0; ledger shows reward 0; UTXO set unchanged."],
        ["Unknown student ID", "Mark with ALU999", "ERROR: Student ID not found"],
        ["Insufficient balance (account)", "Transfer 100 from ALU001 (balance 18)", "Rejected: insufficient account balance."],
        ["Reused / wrong nonce", "Transfer with nonce = 5 when current is 0", "Rejected: incorrect or reused nonce."],
        ["Insufficient UTXOs", "Transfer 100 from ALU001 with 9 unspent", "Rejected: insufficient UTXOs for amount plus fee."],
        ["Unprofitable cloud rental", "Rounds 3, rental fee 22, reward 18", "Warning: cloud rental is unprofitable at round 1, 2, 3."],
        ["Chain tampering", "Mine three blocks, run option 14 on block 1", "Block 1 TAMPERED (hash + signature). Reverted -> VALID."],
    ],
)

# ---------- 10. Design Choices ----------
H1("10. Design Choices and Assumptions")
for line in [
    "Self-contained SHA-256 is bundled in the source so there is no OpenSSL dependency for grading. The same SHA-256 backs the tx_id, block hash, signature, and PoW.",
    "HMAC-style signature is used in place of full ECDSA to keep the Formative 2 binary single-file and dependency-free. The signature still detects field tampering and depends on a system secret, so it satisfies the rubric's requirement that integrity be enforced cryptographically.",
    "In-memory state: the pending pool, confirmed chain, UTXO set, and account history all live in memory for the session. The Formative 1 codebase handles on-disk persistence; this formative focuses on the mining and transaction semantics.",
    "Tail-append linked lists for the pending pool and confirmed chain so iteration order matches the order events occurred.",
    "previous_hash chaining is enforced during mining; the first block's previous_hash is the genesis hash (64 zeros).",
    "Pool size is hardcoded at 3 miners with random hash rates each round.",
]:
    doc.add_paragraph(line, style="List Bullet")

# ---------- 11. Build & Run ----------
H1("11. Build and Run")
code(
    "$ make\n"
    "gcc -Wall -Wextra -Werror -pedantic -std=c99 attendance_mining.c -o attendance_mining\n"
    "\n"
    "$ ./attendance_mining --model utxo --difficulty 3\n"
    "Loaded 5 student(s) from students.txt\n"
    "Model=UTXO difficulty=3 genesis=0000000000000000...\n"
    "...menu..."
)

doc.save("Technical_Report.docx")
print("Technical_Report.docx generated.")
